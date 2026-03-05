"""
File processing module for documents, code, data files, and images.
"""
import csv
import io
import json
import mimetypes
import os
import time
from pathlib import Path
from typing import Dict, Any, List, Tuple
from uuid import uuid4

from PyPDF2 import PdfReader
from docx import Document
from PIL import Image, ExifTags

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False

try:
    import xlrd
    XLRD_AVAILABLE = True
except Exception:
    XLRD_AVAILABLE = False

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except Exception:
    PYTESSERACT_AVAILABLE = False

from config import settings

# Ensure upload directory exists
settings.UPLOAD_DIR.mkdir(exist_ok=True)

TEXT_FILE_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst",
    ".py", ".java", ".js", ".jsx", ".ts", ".tsx",
    ".html", ".htm", ".css", ".scss", ".sass",
    ".yaml", ".yml", ".xml", ".ini", ".cfg", ".conf", ".toml",
    ".sql", ".log", ".sh", ".bat", ".ps1"
}

JSON_EXTENSIONS = {".json"}
CSV_EXTENSIONS = {".csv"}
DOCX_EXTENSIONS = {".docx"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
LEGACY_EXCEL_EXTENSIONS = {".xls"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif"}


def get_file_extension(filename: str) -> str:
    """Get file extension from filename"""
    return Path(filename).suffix.lower()


def is_allowed_file(filename: str) -> bool:
    """Check if file type is allowed"""
    ext = get_file_extension(filename)
    return ext in settings.ALLOWED_EXTENSIONS


def detect_mime_type(file_content: bytes, filename: str) -> str:
    """Detect MIME type from file content and filename"""
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type:
        return mime_type

    if file_content.startswith(b"%PDF"):
        return "application/pdf"
    if file_content.startswith(b"\x89PNG"):
        return "image/png"
    if file_content.startswith(b"\xff\xd8"):
        return "image/jpeg"
    if file_content.startswith(b"GIF8"):
        return "image/gif"
    if file_content.startswith(b"PK"):
        return "application/zip"

    return "application/octet-stream"


def decode_text_bytes(file_content: bytes) -> str:
    """Decode text bytes with tolerant fallbacks."""
    for encoding in ("utf-8", "utf-8-sig", "utf-16", "latin-1"):
        try:
            return file_content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return ""


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from plain text/code/config files"""
    content = decode_text_bytes(file_content)
    if content == "":
        return "[Error: Could not decode text file]"
    return content


def extract_text_from_json(file_content: bytes) -> str:
    """Extract pretty JSON text when possible"""
    raw = decode_text_bytes(file_content)
    if not raw.strip():
        return "[No content found in JSON file]"

    try:
        parsed = json.loads(raw)
        return json.dumps(parsed, indent=2, ensure_ascii=False)
    except Exception:
        # Keep raw content if JSON is not perfectly valid.
        return raw


def extract_text_from_csv(file_content: bytes, max_rows: int = 250) -> str:
    """Extract CSV as plain tabular text"""
    raw = decode_text_bytes(file_content)
    if not raw.strip():
        return "[No content found in CSV file]"

    try:
        reader = csv.reader(io.StringIO(raw))
        lines = []
        for idx, row in enumerate(reader):
            if idx >= max_rows:
                lines.append(f"[CSV truncated to first {max_rows} rows]")
                break
            lines.append(" | ".join(str(cell).strip() for cell in row))
        return "\n".join(lines) if lines else "[No rows found in CSV file]"
    except Exception as exc:
        return f"[Error reading CSV: {exc}]"


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(io.BytesIO(file_content))
        text_parts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                page_text = page_text.strip()
                if page_text:
                    text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            except Exception as exc:
                text_parts.append(f"[Error extracting page {i + 1}: {exc}]")

        if text_parts:
            return "\n\n".join(text_parts)
        return (
            "[No extractable text found in PDF. "
            "This PDF may be scanned/image-based, so OCR may be required.]"
        )
    except Exception as exc:
        return f"[Error reading PDF: {exc}]"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    text_parts.append(" | ".join(row_cells))

        return "\n\n".join(text_parts) if text_parts else "[No text content found in DOCX]"
    except Exception as exc:
        return f"[Error reading DOCX: {exc}]"


def extract_text_from_excel_xlsx(file_content: bytes, max_rows_per_sheet: int = 250) -> str:
    """Extract text from modern Excel files (.xlsx/.xlsm)"""
    if not OPENPYXL_AVAILABLE:
        return "[Error: openpyxl not installed. Install openpyxl to read .xlsx files.]"

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True, read_only=True)
        parts = []
        for sheet in workbook.worksheets:
            parts.append(f"--- Sheet: {sheet.title} ---")
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                row_count += 1
                if row_count > max_rows_per_sheet:
                    parts.append(f"[Sheet truncated to first {max_rows_per_sheet} rows]")
                    break
                values = ["" if value is None else str(value) for value in row]
                if any(v.strip() for v in values):
                    parts.append(" | ".join(values))
        return "\n".join(parts) if parts else "[No content found in Excel workbook]"
    except Exception as exc:
        return f"[Error reading Excel workbook: {exc}]"


def extract_text_from_excel_xls(file_content: bytes, max_rows_per_sheet: int = 250) -> str:
    """Extract text from legacy Excel .xls files"""
    if not XLRD_AVAILABLE:
        return (
            "[Error: xlrd not installed. Legacy .xls is not available in this setup. "
            "Please convert .xls to .xlsx for best support.]"
        )

    try:
        workbook = xlrd.open_workbook(file_contents=file_content)
        parts = []
        for sheet_idx in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_idx)
            parts.append(f"--- Sheet: {sheet.name} ---")
            for row_idx in range(min(sheet.nrows, max_rows_per_sheet)):
                values = [str(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)]
                if any(v.strip() for v in values):
                    parts.append(" | ".join(values))
            if sheet.nrows > max_rows_per_sheet:
                parts.append(f"[Sheet truncated to first {max_rows_per_sheet} rows]")
        return "\n".join(parts) if parts else "[No content found in legacy Excel file]"
    except Exception as exc:
        return f"[Error reading legacy Excel file: {exc}]"


def extract_image_info(file_content: bytes, filename: str) -> str:
    """Extract image metadata and OCR text (if OCR is available)."""
    try:
        image_file = io.BytesIO(file_content)
        img = Image.open(image_file)

        info_parts = [
            f"Image: {filename}",
            f"Format: {img.format}",
            f"Size: {img.size[0]}x{img.size[1]} pixels",
            f"Mode: {img.mode}",
        ]

        # EXIF metadata
        try:
            exif = img._getexif()  # pylint: disable=protected-access
            if exif:
                exif_data = []
                for tag_id, value in exif.items():
                    tag = ExifTags.TAGS.get(tag_id, tag_id)
                    if tag in {"DateTime", "Make", "Model", "LensModel", "Software"}:
                        exif_data.append(f"{tag}: {value}")
                if exif_data:
                    info_parts.append("EXIF: " + ", ".join(exif_data))
        except Exception:
            pass

        ocr_text = ""
        if PYTESSERACT_AVAILABLE:
            try:
                ocr_text = pytesseract.image_to_string(img).strip()
            except Exception as exc:
                info_parts.append(f"[OCR unavailable at runtime: {exc}]")
        else:
            info_parts.append("[OCR library not installed. Install pytesseract + Tesseract OCR for image text extraction.]")

        if ocr_text:
            info_parts.append("\nOCR Text:\n" + ocr_text)
        else:
            info_parts.append("\n[No OCR text extracted from image.]")

        return "\n".join(info_parts)
    except Exception as exc:
        return f"[Error reading image: {exc}]"


def process_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """Process a file and extract usable text content."""
    mime_type = detect_mime_type(file_content, filename)
    ext = get_file_extension(filename)

    if len(file_content) > settings.MAX_FILE_SIZE:
        return {
            "filename": filename,
            "mime_type": mime_type,
            "size": len(file_content),
            "content": f"[Error: File exceeds maximum size of {settings.MAX_FILE_SIZE / (1024 * 1024):.1f}MB]",
            "success": False,
        }

    if not is_allowed_file(filename):
        return {
            "filename": filename,
            "mime_type": mime_type,
            "size": len(file_content),
            "content": f"[Error: File type '{ext}' is not supported. Allowed types: {', '.join(sorted(settings.ALLOWED_EXTENSIONS))}]",
            "success": False,
        }

    success = True
    if ext in TEXT_FILE_EXTENSIONS:
        content = extract_text_from_txt(file_content)
    elif ext in JSON_EXTENSIONS:
        content = extract_text_from_json(file_content)
    elif ext in CSV_EXTENSIONS:
        content = extract_text_from_csv(file_content)
    elif ext in PDF_EXTENSIONS:
        content = extract_text_from_pdf(file_content)
    elif ext in DOCX_EXTENSIONS:
        content = extract_text_from_docx(file_content)
    elif ext in EXCEL_EXTENSIONS:
        content = extract_text_from_excel_xlsx(file_content)
    elif ext in LEGACY_EXCEL_EXTENSIONS:
        content = extract_text_from_excel_xls(file_content)
    elif ext in IMAGE_EXTENSIONS:
        content = extract_image_info(file_content, filename)
    elif mime_type.startswith("text/"):
        content = extract_text_from_txt(file_content)
    else:
        content = f"[Error: Unsupported file type '{mime_type}' ({ext})]"
        success = False

    if content.startswith("[Error:"):
        success = False

    max_chars = 80000
    if len(content) > max_chars:
        original_len = len(content)
        content = content[:max_chars] + f"\n\n[Content truncated. Original length: {original_len} characters]"

    return {
        "filename": filename,
        "mime_type": mime_type,
        "size": len(file_content),
        "content": content,
        "success": success,
    }


def process_multiple_files(files: List[Tuple[bytes, str]]) -> List[Dict[str, Any]]:
    """Process multiple files"""
    return [process_file(file_content, filename) for file_content, filename in files]


def format_files_for_prompt(file_results: List[Dict[str, Any]]) -> str:
    """Format processed files for inclusion in the model prompt"""
    if not file_results:
        return ""

    formatted_parts = [
        "=" * 50,
        "ATTACHED FILES",
        "=" * 50,
    ]

    for i, file_result in enumerate(file_results, 1):
        formatted_parts.append(f"\n--- File {i}: {file_result['filename']} ---")
        formatted_parts.append(f"Type: {file_result['mime_type']}")
        formatted_parts.append(f"Size: {file_result['size']} bytes")
        formatted_parts.append(f"Status: {'Success' if file_result['success'] else 'Failed'}")
        formatted_parts.append("\nContent:")
        formatted_parts.append(file_result["content"])
        formatted_parts.append("-" * 40)

    return "\n".join(formatted_parts)


def save_uploaded_file(file_content: bytes, filename: str) -> Path:
    """Save uploaded file to disk with a unique, stable filename."""
    original_name = Path(filename).name
    stem = Path(original_name).stem or "upload"
    suffix = Path(original_name).suffix
    timestamp = int(time.time() * 1000)
    unique_id = uuid4().hex[:8]
    safe_filename = f"{timestamp}_{unique_id}_{stem}{suffix}"
    file_path = settings.UPLOAD_DIR / safe_filename

    with open(file_path, "wb") as file_obj:
        file_obj.write(file_content)

    return file_path


def cleanup_old_uploads(max_age_hours: int = 24):
    """Clean up old uploaded files"""
    import time

    current_time = time.time()
    max_age_seconds = max_age_hours * 3600

    for file_path in settings.UPLOAD_DIR.iterdir():
        if not file_path.is_file():
            continue
        file_age = current_time - os.path.getmtime(file_path)
        if file_age > max_age_seconds:
            try:
                file_path.unlink()
                print(f"Deleted old upload: {file_path.name}")
            except Exception as exc:
                print(f"Error deleting {file_path}: {exc}")
